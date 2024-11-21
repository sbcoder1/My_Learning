from docx import Document
from docx.shared import Pt
import mysql.connector
from datetime import datetime

def myRecord():
    
    db_Conn = mysql.connector.connect(
        host="localhost",
        user="root",
        password="root",
        database="systeminfo"  
    )

    cursor = db_Conn.cursor()
    cursor = db_Conn.cursor()
    cursor = db_Conn.cursor()
    cursor = db_Conn.cursor()
    query1 = """SELECT version,system,sys_release,uname,n_node,machine,processor,architecture
               FROM systeminfo.platforminfo;"""
    query2 = """SELECT pysical_cores,max_freq,min_freq,cpu_usage,cpu_total
               FROM systeminfo.cpuinfo;"""
    query3 = """SELECT file_type,total_memory,avaible,used,percentage
               FROM systeminfo.memoryinfo;"""
    query4 = """SELECT hostname,ip_address,byte_send,byte_receive
               FROM systeminfo.networkinfo;"""

    cursor.execute(query1)
    getrec1 = cursor.fetchall()
    cursor.execute(query2)
    getrec2 = cursor.fetchall()
    cursor.execute(query3)
    getrec3 = cursor.fetchall()
    cursor.execute(query4)
    getrec4 = cursor.fetchall()

    print(getrec1)
    print(getrec2)
    print(getrec3)
    print(getrec4)
    

    d = Document()

    section = d.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width  
    section.page_height = new_height  


    h =d.add_heading("System Information", 0)
    h_run = h.runs[0]
    h_run.font.size = Pt(30)   
    h_run.bold = True

    c_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    d.add_paragraph(f"System Info Data Fetch On: {c_datetime}")

    p = d.add_paragraph('Platform Info ')
    p_run = p.runs[0]
    p_run.font.size = Pt(18)   
    p_run.bold = True
    
    table = d.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'version'
    hdr_cells[1].text = 'system'
    hdr_cells[2].text = 'release'
    hdr_cells[3].text = 'uname'
    hdr_cells[4].text = 'n_node'
    hdr_cells[5].text = 'machine'
    hdr_cells[6].text = 'processor'
    hdr_cells[7].text = 'architecture'

    for version,system,release,uname,n_node,machine,processor,architecture in getrec1:
        row_cells = table.add_row().cells  
        row_cells[0].text = str(version)
        row_cells[1].text = str(system)
        row_cells[2].text = str(release)
        row_cells[3].text = str(uname)
        row_cells[4].text = str(n_node)
        row_cells[5].text = str(machine)
        row_cells[6].text = str(processor)
        row_cells[7].text = str(architecture)
    #------------------------------------------------------------------

    p = d.add_paragraph('Cpu Info ')
    p_run = p.runs[0]
    p_run.font.size = Pt(18)   
    p_run.bold = True    
    
    table = d.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'pysical_cores'
    hdr_cells[1].text = 'max_freq'
    hdr_cells[2].text = 'min_freq'
    hdr_cells[3].text = 'cpu_usage'
    hdr_cells[4].text = 'cpu_total'


    for pysical_cores,max_freq,min_freq,cpu_usage,cpu_total in getrec2:
        row_cells = table.add_row().cells  
        row_cells[0].text = str(pysical_cores)
        row_cells[1].text = str(max_freq)
        row_cells[2].text = str(min_freq)
        row_cells[3].text = str(cpu_usage)
        row_cells[4].text = str(cpu_total)


    #------------------------------------------------
    p = d.add_paragraph('Memory Info ')
    p_run = p.runs[0]
    p_run.font.size = Pt(18)   
    p_run.bold = True
    
    table = d.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'file_type'
    hdr_cells[1].text = 'total_memory'
    hdr_cells[2].text = 'avaible'
    hdr_cells[3].text = 'used'
    hdr_cells[4].text = 'percentage'




    for file_type,total_memory,avaible,used,percentage in getrec3:
        row_cells = table.add_row().cells  
        row_cells[0].text = str(file_type)
        row_cells[1].text = str(total_memory)
        row_cells[2].text = str(avaible)
        row_cells[3].text = str(used)
        row_cells[4].text = str(percentage)
       
        


    #--------------------------------------
    p = d.add_paragraph('NetWork Info ')
    p_run = p.runs[0]
    p_run.font.size = Pt(18)   
    p_run.bold = True

    table = d.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'hostname'
    hdr_cells[1].text = 'ip_address'
    hdr_cells[2].text = 'byte_send'
    hdr_cells[3].text = 'byte_receive'





    for hostname,ip_address,byte_send,byte_receive in getrec4:
        row_cells = table.add_row().cells  
        row_cells[0].text = str(hostname)
        row_cells[1].text = str(ip_address)
        row_cells[2].text = str(byte_send)
        row_cells[3].text = str(byte_receive)
       
    c_datetime = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    d.save(f"C:\\Users\\sbcod\\Desktop\\Study Files\\SystemInfo_{c_datetime}.docx")

    cursor.close()
    db_Conn.close()
    
myRecord()
print("Today System Info has been saved successfully")   

