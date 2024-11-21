from docx import Document
from docx.shared import Pt
import mysql.connector
from datetime import datetime


db_Conn = mysql.connector.connect(
    host="localhost",
    user="root",
    password="root",
    database="weather"  
)

cursor = db_Conn.cursor()

query = """SELECT sys_country, location_name, weather_main, temp, humidity, timezone
           FROM weather_data;"""  

cursor.execute(query)  
getrec = cursor.fetchall()  

if not getrec:
    print("No records found in the weather_data table.")
    exit(0)  

d = Document()

h = d.add_heading("Weather Report", 0)
h_run = h.runs[0]
h_run.font.size = Pt(30)   
h_run.bold = True

c_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
d.add_paragraph(f"Weather Data Fetch On: {c_datetime}")


table = d.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'sys_country'
hdr_cells[1].text = 'location_name'
hdr_cells[2].text = 'weather_main'
hdr_cells[3].text = 'temp'
hdr_cells[4].text = 'humidity'
hdr_cells[5].text = 'timezone'



for sys_country, location_name, weather_main, temp, humidity, timezone in getrec:
    row_cells = table.add_row().cells  
    row_cells[0].text = str(sys_country)
    row_cells[1].text = str(location_name)
    row_cells[2].text = str(weather_main)
    row_cells[3].text = str(temp)
    row_cells[4].text = str(humidity)
    row_cells[5].text = str(timezone)
    

d.save("C:\\Users\\sbcod\\Desktop\\Study Files\\22102024-7PM.docx")

cursor.close()
db_Conn.close()

print("Today Weather report has been saved successfully")
